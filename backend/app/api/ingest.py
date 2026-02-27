"""Document ingestion endpoint.

Only admin users can ingest documents (RBAC).
Supports: .md, .txt, .pdf, .docx
"""

from __future__ import annotations

import hashlib
import io

from fastapi import APIRouter, Depends, HTTPException, UploadFile, status
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.auth.dependencies import require_admin
from app.models.database import get_db
from app.models.db import Document, DocumentChunk, User
from app.observability import get_logger
from app.rag.chunker import chunk_document
from app.rag.embedder import embed_texts
from app.rag.vector_store import get_vector_store

log = get_logger(__name__)
router = APIRouter(prefix="/api/ingest", tags=["ingest"])

ALLOWED_MIME = {
    "text/plain", "text/markdown",
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB


class IngestResponse(BaseModel):
    document_id: int
    filename: str
    chunks_created: int
    deduplicated: bool


def _extract_text(filename: str, content: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext in ("md", "txt"):
        return content.decode("utf-8", errors="replace")
    elif ext == "pdf":
        import pypdf  # noqa: PLC0415

        reader = pypdf.PdfReader(io.BytesIO(content))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    elif ext == "docx":
        import docx  # noqa: PLC0415

        doc = docx.Document(io.BytesIO(content))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        raise ValueError(f"Unsupported file type: {ext}")


@router.post("", response_model=IngestResponse, status_code=status.HTTP_201_CREATED)
async def ingest_document(
    file: UploadFile,
    admin: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="File too large (max 10 MB)")

    content_hash = hashlib.sha256(content).hexdigest()
    filename = file.filename or "upload"

    # Dedup check
    existing = await db.execute(select(Document).where(Document.content_hash == content_hash))
    if existing.scalar_one_or_none():
        log.info("ingest_deduplicated", filename=filename, hash=content_hash[:12])
        doc = existing.scalar_one_or_none()
        return IngestResponse(
            document_id=doc.id, filename=filename, chunks_created=0, deduplicated=True
        )

    # Parse text
    try:
        text = _extract_text(filename, content)
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Failed to parse file: {exc}") from exc

    if not text.strip():
        raise HTTPException(status_code=422, detail="Document appears to be empty after parsing")

    # Chunk
    chunks = chunk_document(text, filename)
    if not chunks:
        raise HTTPException(status_code=422, detail="No usable chunks extracted from document")

    # Save document record
    doc = Document(
        filename=filename,
        content_hash=content_hash,
        chunk_count=len(chunks),
        uploaded_by=admin.id,
    )
    db.add(doc)
    await db.flush()  # get doc.id

    # Embed all chunks in one batch
    texts = [c.text for c in chunks]
    vectors = embed_texts(texts)

    # Save chunks + index into vector store
    store = get_vector_store()
    chunk_records: list[DocumentChunk] = []
    for idx, (c, vec) in enumerate(zip(chunks, vectors)):
        rec = DocumentChunk(
            document_id=doc.id,
            chunk_index=idx,
            text=c.text,
            source_label=c.source_label,
        )
        db.add(rec)
        chunk_records.append(rec)

    await db.flush()  # get chunk IDs

    for rec, vec in zip(chunk_records, vectors):
        await store.add(
            chunk_id=rec.id,
            text=rec.text,
            source_label=rec.source_label,
            vector=vec,
        )

    await db.commit()
    await store.flush()

    log.info("ingest_complete", document_id=doc.id, filename=filename, chunks=len(chunks))
    return IngestResponse(document_id=doc.id, filename=filename, chunks_created=len(chunks), deduplicated=False)


@router.delete("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(
    document_id: int,
    admin: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc: Document | None = result.scalar_one_or_none()
    if doc is None:
        raise HTTPException(status_code=404, detail="Document not found")

    chunk_ids = [c.id for c in doc.chunks]
    store = get_vector_store()
    await store.delete_by_document(chunk_ids)
    await store.flush()

    await db.delete(doc)
    await db.commit()


@router.get("", response_model=list[dict])
async def list_documents(
    admin: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(Document))
    docs = result.scalars().all()
    return [
        {
            "id": d.id,
            "filename": d.filename,
            "chunk_count": d.chunk_count,
            "created_at": d.created_at.isoformat(),
        }
        for d in docs
    ]
